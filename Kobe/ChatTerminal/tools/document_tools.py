"""
文档提取工具模块
支持 PDF, DOC, DOCX, 图片OCR等格式
"""

import os
from pathlib import Path
from typing import Optional
import mimetypes


def extract_document(file_path: str, max_length: int = 50000) -> str:
    """
    从文档中提取文本内容
    支持格式: PDF, DOC, DOCX, TXT, 图片(OCR)
    
    Args:
        file_path: 文件路径（可以是绝对路径或相对于downloads/的路径）
        max_length: 最大返回字符数（避免超过token限制）
    
    Returns:
        提取的文本内容
    """
    try:
        # 处理路径
        path = Path(file_path)
        
        # 如果是相对路径，尝试在downloads目录查找
        if not path.is_absolute():
            downloads_dir = Path(__file__).parent.parent / "downloads"
            path = downloads_dir / file_path
        
        # 检查文件是否存在
        if not path.exists():
            return f"错误: 文件不存在 - {path}"
        
        # 获取文件类型
        file_ext = path.suffix.lower()
        
        # 根据文件类型提取内容
        if file_ext == '.pdf':
            content = _extract_pdf(path)
        elif file_ext in ['.doc', '.docx']:
            content = _extract_word(path)
        elif file_ext in ['.txt', '.md', '.log', '.json', '.xml', '.csv']:
            content = _extract_text(path)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif']:
            content = _extract_image_ocr(path)
        else:
            return f"错误: 不支持的文件格式 '{file_ext}'\n支持的格式: PDF, DOC, DOCX, TXT, PNG, JPG等"
        
        # 限制长度
        if len(content) > max_length:
            content = content[:max_length] + f"\n\n... (内容过长，已截断。原文共 {len(content)} 字符)"
        
        return content
    
    except Exception as e:
        return f"错误: 提取文档内容时出错 - {str(e)}"


def _extract_pdf(file_path: Path) -> str:
    """提取PDF文本"""
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- 第 {page_num} 页 ---\n{text}\n")
            
            if not text_parts:
                return "警告: PDF文件为空或无法提取文本（可能是扫描版PDF，需要OCR）"
            
            return "\n".join(text_parts)
    
    except ImportError:
        return "错误: 缺少 PyPDF2 库。请运行: pip install PyPDF2"
    except Exception as e:
        return f"错误: 提取PDF时出错 - {str(e)}"


def _extract_word(file_path: Path) -> str:
    """提取Word文档文本"""
    try:
        import docx
        
        doc = docx.Document(file_path)
        text_parts = []
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        if not text_parts:
            return "警告: Word文档为空"
        
        return "\n\n".join(text_parts)
    
    except ImportError:
        return "错误: 缺少 python-docx 库。请运行: pip install python-docx"
    except Exception as e:
        return f"错误: 提取Word文档时出错 - {str(e)}"


def _extract_text(file_path: Path) -> str:
    """提取纯文本文件"""
    try:
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    return content
            except UnicodeDecodeError:
                continue
        
        return "错误: 无法识别文件编码"
    
    except Exception as e:
        return f"错误: 读取文本文件时出错 - {str(e)}"


def _extract_image_ocr(file_path: Path) -> str:
    """使用OCR提取图片中的文本"""
    try:
        from PIL import Image
        import pytesseract
        
        # 打开图片
        image = Image.open(file_path)
        
        # OCR识别（支持中英文）
        text = pytesseract.image_to_string(image, lang='chi_sim+eng')
        
        if not text.strip():
            return "警告: 图片中未识别到文本"
        
        return f"--- OCR识别结果 ---\n{text}"
    
    except ImportError as e:
        if 'PIL' in str(e):
            return "错误: 缺少 Pillow 库。请运行: pip install Pillow"
        elif 'pytesseract' in str(e):
            return "错误: 缺少 pytesseract 库。请运行: pip install pytesseract\n并安装 Tesseract-OCR: https://github.com/tesseract-ocr/tesseract"
        else:
            return f"错误: 缺少必要的库 - {str(e)}"
    except Exception as e:
        return f"错误: OCR识别时出错 - {str(e)}"


if __name__ == "__main__":
    # 测试
    import sys
    if len(sys.argv) > 1:
        print(extract_document(sys.argv[1]))
    else:
        print("用法: python document_tools.py <文件路径>")

